import re
import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


LINKS = [
    "https://inter.kubsu.ru/portfolio/faculty-of-computer-technology-and-applied-mathematics/",
    "https://inter.kubsu.ru/portfolio/faculty-of-biology/",
    "https://inter.kubsu.ru/portfolio/faculty-of-art-and-graphics/",
    "https://inter.kubsu.ru/portfolio/faculty-of-chemistry-and-high-technologies/",
    "https://inter.kubsu.ru/portfolio/faculty-of-architecture-and-design/"
    # добавь сюда свои ссылки
]

OUTPUT_FILE = "faculties.docx"

CATEGORY_SUFFIXES = {
    "about": "About",
    "dean": "Dean",
    "bach": "Bachelor programmes",
    "mag": "Master programmes",
    "depart": "Departments",
    "science": "Science",
    "inter": "International activity",
    "contacts": "Contacts",
}


def clean_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def fetch_html(url: str, timeout: int = 30) -> str:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0 Safari/537.36"
        )
    }
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()
    return response.text


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        text = clean_text(item)
        if text:
            document.add_paragraph(text, style="List Bullet")


def add_numbered_list(document: Document, items: list[str]) -> None:
    for item in items:
        text = clean_text(item)
        if text:
            document.add_paragraph(text, style="List Number")


def extract_block_content(document: Document, block: Tag) -> None:
    """
    Универсально извлекает текст из блока:
    - p
    - ul / ol
    - h1..h6
    - table
    - прочие теги
    """
    for child in block.children:
        if not isinstance(child, Tag):
            continue

        if child.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                level = 3 if child.name in {"h1", "h2"} else 4
                document.add_heading(text, level=level)

        elif child.name == "p":
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                document.add_paragraph(text)

        elif child.name == "ul":
            items = [li.get_text(" ", strip=True) for li in child.find_all("li", recursive=False)]
            add_bullet_list(document, items)

        elif child.name == "ol":
            items = [li.get_text(" ", strip=True) for li in child.find_all("li", recursive=False)]
            add_numbered_list(document, items)

        elif child.name == "table":
            rows = child.find_all("tr")
            for row in rows:
                cells = row.find_all(["th", "td"])
                cell_texts = [clean_text(c.get_text(" ", strip=True)) for c in cells]
                cell_texts = [t for t in cell_texts if t]
                if cell_texts:
                    document.add_paragraph(" | ".join(cell_texts))

        elif child.name in {"div", "section", "article"}:
            extract_block_content(document, child)

        else:
            text = clean_text(child.get_text(" ", strip=True))
            if text:
                document.add_paragraph(text)


def parse_spoilers(document: Document, soup: BeautifulSoup) -> None:
    accordion = soup.select_one("div.su-accordion")
    if not accordion:
        return

    spoilers = accordion.select("div.su-spoiler")
    if not spoilers:
        return

    document.add_heading("Accordion sections", level=2)

    for spoiler in spoilers:
        title_tag = spoiler.select_one("div.su-spoiler-title")
        title = clean_text(title_tag.get_text(" ", strip=True)) if title_tag else "Untitled section"
        document.add_heading(title, level=3)

        content = spoiler.select_one("div.su-spoiler-content")
        if content:
            extract_block_content(document, content)
        else:
            document.add_paragraph("No content found.")

        document.add_paragraph("")


def find_category_blocks(soup: BeautifulSoup) -> list[Tag]:
    """
    Ищет любые элементы с id, оканчивающимся на:
    _about, _dean, _bach, _mag, _depart, _science, _inter, _contacts
    """
    suffix_pattern = "|".join(re.escape(suffix) for suffix in CATEGORY_SUFFIXES.keys())
    pattern = re.compile(rf".*_({suffix_pattern})$")

    found = []
    for tag in soup.find_all(attrs={"id": True}):
        tag_id = tag.get("id", "")
        if pattern.fullmatch(tag_id):
            found.append(tag)

    return found


def pretty_section_name_from_id(tag_id: str) -> str:
    """
    Например:
    mmf_about -> mmf / About
    fit_bach -> fit / Bachelor programmes
    """
    for suffix, title in CATEGORY_SUFFIXES.items():
        end = "_" + suffix
        if tag_id.endswith(end):
            faculty_name = tag_id[:-len(end)]
            faculty_name = faculty_name.strip("_- ")
            return f"{faculty_name} — {title}" if faculty_name else title
    return tag_id


def parse_category_blocks(document: Document, soup: BeautifulSoup) -> None:
    blocks = find_category_blocks(soup)
    if not blocks:
        return

    document.add_heading("Faculty sections by id", level=2)

    used_ids = set()

    for block in blocks:
        tag_id = block.get("id", "").strip()
        if not tag_id or tag_id in used_ids:
            continue
        used_ids.add(tag_id)

        section_title = pretty_section_name_from_id(tag_id)
        document.add_heading(section_title, level=3)

        extract_block_content(document, block)
        document.add_paragraph("")


def parse_page_to_doc(document: Document, soup: BeautifulSoup, source_url: str) -> None:
    page_title = clean_text(soup.title.get_text(" ", strip=True)) if soup.title else source_url
    document.add_heading(page_title, level=1)

    p = document.add_paragraph()
    run = p.add_run(f"Source: {source_url}")
    run.italic = True

    parse_spoilers(document, soup)
    parse_category_blocks(document, soup)


def build_docx_from_links(links: list[str], output_file: str) -> None:
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = document.add_heading("Collected faculty data", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, url in enumerate(links, start=1):
        try:
            html = fetch_html(url)
            soup = BeautifulSoup(html, "lxml")

            if i > 1:
                document.add_page_break()

            parse_page_to_doc(document, soup, url)

        except Exception as e:
            if i > 1:
                document.add_page_break()
            document.add_heading("Processing error", level=1)
            document.add_paragraph(url)
            document.add_paragraph(str(e))

    document.save(output_file)
    print(f"Done: {output_file}")


if __name__ == "__main__":
    build_docx_from_links(LINKS, OUTPUT_FILE)