import requests
from bs4 import BeautifulSoup, Tag
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


LINKS = [
    "https://inter.kubsu.ru/portfolio/faculty-of-computer-technology-and-applied-mathematics/",
    "https://inter.kubsu.ru/portfolio/faculty-of-biology/",
    "https://inter.kubsu.ru/portfolio/faculty-of-art-and-graphics/",
    "https://inter.kubsu.ru/portfolio/faculty-of-chemistry-and-high-technologies/",
    "https://inter.kubsu.ru/portfolio/faculty-of-architecture-and-design/"
    # добавь сюда свои ссылки
]

OUTPUT_FILE = "programmes.docx"


def clean_text(text: str) -> str:
    return " ".join(text.replace("\xa0", " ").split()).strip()


def fetch_html(url: str, timeout: int = 30) -> str:
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0 Safari/537.36"
        )
    }
    response = requests.get(url, headers=headers, timeout=timeout)
    response.raise_for_status()
    return response.text


def add_bullet_list(document: Document, items: list[str]) -> None:
    for item in items:
        item = clean_text(item)
        if item:
            document.add_paragraph(item, style="List Bullet")


def parse_spoilers_to_doc(document: Document, soup: BeautifulSoup, source_url: str) -> None:
    # Заголовок страницы
    page_title = clean_text(soup.title.get_text(" ", strip=True)) if soup.title else source_url
    document.add_heading(page_title, level=1)

    # Ссылка-источник
    p = document.add_paragraph()
    run = p.add_run(f"Источник: {source_url}")
    run.italic = True

    accordion = soup.select_one("div.su-accordion")
    if not accordion:
        document.add_paragraph("Блок su-accordion не найден.")
        return

    spoilers = accordion.select("div.su-spoiler")
    if not spoilers:
        document.add_paragraph("Блоки su-spoiler не найдены.")
        return

    for spoiler in spoilers:
        title_tag = spoiler.select_one("div.su-spoiler-title")
        title = clean_text(title_tag.get_text(" ", strip=True)) if title_tag else "Без названия"
        document.add_heading(title, level=2)

        content = spoiler.select_one("div.su-spoiler-content")
        if not content:
            document.add_paragraph("Нет содержимого.")
            continue

        # Идем по прямым потомкам контента, чтобы сохранить структуру
        for child in content.children:
            if not isinstance(child, Tag):
                continue

            # Абзацы
            if child.name == "p":
                text = clean_text(child.get_text(" ", strip=True))
                if text:
                    document.add_paragraph(text)

            # Списки
            elif child.name in ("ul", "ol"):
                items = [li.get_text(" ", strip=True) for li in child.find_all("li", recursive=False)]
                add_bullet_list(document, items)

            # На случай других тегов
            else:
                text = clean_text(child.get_text(" ", strip=True))
                if text:
                    document.add_paragraph(text)

        document.add_paragraph("")  # пустая строка между программами


def build_docx_from_links(links: list[str], output_file: str) -> None:
    document = Document()

    # Небольшое оформление
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    title = document.add_heading("Собранные данные по программам", level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    for i, url in enumerate(links, start=1):
        try:
            html = fetch_html(url)
            soup = BeautifulSoup(html, "lxml")

            if i > 1:
                document.add_page_break()

            parse_spoilers_to_doc(document, soup, url)

        except Exception as e:
            document.add_heading(f"Ошибка при обработке: {url}", level=1)
            document.add_paragraph(str(e))

    document.save(output_file)
    print(f"Готово: {output_file}")


if __name__ == "__main__":
    build_docx_from_links(LINKS, OUTPUT_FILE)